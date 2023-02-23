import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Cm
import math

def merge(doc_final):
    #prepara o doc final para o merge
    composer = Composer(doc_final)
    #abre o doc temporário
    doc_merge = Document('docx_rendered.docx')
    #faz merge e salva o doc final
    composer.append(doc_merge)
    composer.save("declaracao.docx")

def merge_dest(doc_final):
    #prepara o doc final para o merge
    composer = Composer(doc_final)
    #abre o doc temporário
    doc_merge = Document('destinatarios_rendered.docx')
    #faz merge e salva o doc finals
    composer.append(doc_merge)
    composer.save("destinatarios_camiseta.docx")

def monta_declaracao():
    try:
        #abre um documento final
        doc_final = Document()
        doc_final.save('declaracao.docx')
        #abre o template
        doc = DocxTemplate('docx_template.docx')
        #abre a planilha de dados 
        xls = pd.ExcelFile("xls_data.xlsx")

        #pega o numero de linhas 
        df = pd.read_excel(xls)
        n = len(df.index)
        #enquanto houverem linhas não tratadas na panilha 
        i = 1
        while (i <= n):
            #le uma linha 
            df = pd.read_excel(xls, header = 0, skiprows = [j for j in range(1,i)], nrows=1,  usecols= 'A:H')
            #transforma o dataframe em dicionario com lista
            data = df.to_dict('list')
            #pega a lista de chaves (cabeçalho da planilha) para o contexto do template
            keys = list(data)
            #se a iteração nao for a última
            if (i < n):
                #le outra linha
                i += 1
                df2 = pd.read_excel(xls, header = 0, skiprows = [j for j in range(1,i)], nrows=1,  usecols= 'A:H')
                data2 = df2.to_dict('list')
                #monta o contexto
                context = {
                    keys[0]: data[keys[0]][0],
                    keys[1]: data[keys[1]][0],
                    keys[2]: data[keys[2]][0],
                    keys[3]: data[keys[3]][0],
                    keys[4]: data[keys[4]][0],
                    keys[5]: data[keys[5]][0],
                    keys[6]: data[keys[6]][0],
                    keys[7]: data[keys[7]][0],
                    keys[0]+"_": data2[keys[0]][0],
                    keys[1]+"_": data2[keys[1]][0],
                    keys[2]+"_": data2[keys[2]][0],
                    keys[3]+"_": data2[keys[3]][0],
                    keys[4]+"_": data2[keys[4]][0],
                    keys[5]+"_": data2[keys[5]][0],
                    keys[6]+"_": data2[keys[6]][0],
                    keys[7]+"_": data2[keys[7]][0],
                }
            #se a iteração for a última, não le outra linha
            else :
                context = {
                    keys[0]: data[keys[0]][0],
                    keys[1]: data[keys[1]][0],
                    keys[2]: data[keys[2]][0],
                    keys[3]: data[keys[3]][0],
                    keys[4]: data[keys[4]][0],
                    keys[5]: data[keys[5]][0],
                    keys[6]: data[keys[6]][0],
                    keys[7]: data[keys[7]][0],
                }
            #renderiza e salva um doc temporário a partir do template com o contexto
            doc.render(context)
            doc.save("docx_rendered.docx")
            #faz merge com o doc final
            merge(doc_final)
            i += 1

        #Ajuste das margens do arquivo após o final do processo
        sections = doc_final.sections
        for section in sections:
            section.top_margin = Cm(0.1)
            section.bottom_margin = Cm(0.1)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
        doc_final.save('declaracao.docx')

    except PackageNotFoundError:
        print('Erro ao abrir o arquivo. O arquivo pode estar corrompido.')

def monta_destinatario():
    try:
        doc_final = Document()
        doc_final.save('destinatarios_camiseta.docx')
        doc = DocxTemplate('destinatarios_template.docx')
        xls = pd.ExcelFile("xls_data.xlsx")

        df = pd.read_excel(xls)
        n = len(df.index)

        n_pag = math.floor(n / 10)
        n_rest = n % 10
    
        #enquanto houverem linhas não tratadas na panilha 
        i = 0
        while (i < n_pag):
            df = pd.read_excel(xls, header = 0, skiprows = [j for j in range(1, i*10+1)], nrows=10,  usecols= 'A:H')
            data = df.to_dict('list')
            keys = list(data)
            context_acumulado = {}
            for j in range(10):
                context = {
                    str(keys[0])+str(j): data[keys[0]][j],
                    str(keys[1])+str(j): data[keys[1]][j],
                    str(keys[2])+str(j): data[keys[2]][j],
                    str(keys[3])+str(j): data[keys[3]][j],
                    str(keys[4])+str(j): data[keys[4]][j],
                    str(keys[5])+str(j): data[keys[5]][j],
                    str(keys[6])+str(j): data[keys[6]][j],
                    str(keys[7])+str(j): data[keys[7]][j],
                }
                context_acumulado.update(context)
            doc.render(context_acumulado)
            doc.save("destinatarios_rendered.docx")
            merge_dest(doc_final)
            i += 1
        else :
            df = pd.read_excel(xls, header = 0, skiprows = [j for j in range(1, i*10+1)], nrows = n_rest,  usecols= 'A:H')
            data = df.to_dict('list')
            keys = list(data)
            context_acumulado = {}
            for j in range(n_rest):
                context = {
                    str(keys[0])+str(j): data[keys[0]][j],
                    str(keys[1])+str(j): data[keys[1]][j],
                    str(keys[2])+str(j): data[keys[2]][j],
                    str(keys[3])+str(j): data[keys[3]][j],
                    str(keys[4])+str(j): data[keys[4]][j],
                    str(keys[5])+str(j): data[keys[5]][j],
                    str(keys[6])+str(j): data[keys[6]][j],
                    str(keys[7])+str(j): data[keys[7]][j],
                }
                context_acumulado.update(context)
            doc.render(context_acumulado)
            doc.save("destinatarios_rendered.docx")
            merge_dest(doc_final)

    except PackageNotFoundError:
        print('Erro ao abrir o arquivo. O arquivo pode estar corrompido.')

def main():
    monta_declaracao()
    monta_destinatario()

if __name__ == "__main__":
    main()